"""Text extraction from uploaded files (PDF, Word, plain text/markdown/CSV)."""
from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)


def extract_text(data: bytes, filename: str, mime: str = "") -> str:
    name = filename.lower()
    try:
        if name.endswith(".pdf") or "pdf" in mime:
            return _pdf(data)
        if name.endswith(".docx") or "wordprocessingml" in mime:
            return _docx(data)
        if name.endswith((".txt", ".md", ".csv", ".json", ".html")) or mime.startswith("text/"):
            return data.decode("utf-8", errors="replace")
    except Exception:
        logger.exception("Extraction failed for %s", filename)
    return ""


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part.strip()).strip()


def chunk_document_text(text: str, target: int = 1200, overlap: int = 150) -> list[str]:
    """Paragraph-aware chunking with light overlap for retrieval continuity."""
    text = text.strip()
    if not text:
        return []
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        while len(para) > target:  # single huge paragraph — hard split
            if current:
                chunks.append(current)
                current = ""
            chunks.append(para[:target])
            para = para[max(0, target - overlap):]
        if len(current) + len(para) + 2 > target and current:
            chunks.append(current)
            current = current[-overlap:] if overlap else ""
        current = (current + "\n\n" + para).strip() if current else para
    if current:
        chunks.append(current)
    return chunks
